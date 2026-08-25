import importlib.util
import os
import tempfile
import unittest
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SPEC = importlib.util.spec_from_file_location("qc_module", PROJECT_ROOT / "qc.py")
qc = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(qc)


class ManuscriptLoadingTests(unittest.TestCase):
    def test_load_txt(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", encoding="utf-8", delete=False) as f:
            f.write("测试正文\n10.1000/example")
            path = f.name
        try:
            self.assertIn("测试正文", qc.load_manuscript(path))
        finally:
            os.unlink(path)

    def test_load_docx_paragraphs_and_tables(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            doc = Document()
            doc.add_paragraph("Word 正文")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "字段"
            table.cell(0, 1).text = "内容"
            doc.save(path)

            text = qc.load_manuscript(path)
            self.assertIn("Word 正文", text)
            self.assertIn("字段 | 内容", text)
        finally:
            os.unlink(path)

    def test_rejects_legacy_doc(self):
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
            path = f.name
        try:
            with self.assertRaisesRegex(RuntimeError, "旧版 .doc"):
                qc.load_manuscript(path)
        finally:
            os.unlink(path)


class AgentDefinitionTests(unittest.TestCase):
    SLUGS = ("qc-conductor", "qc-verifier", "qc-reporter")

    def test_every_agent_definition_loads(self):
        for slug in self.SLUGS:
            with self.subTest(slug=slug):
                body = qc.read_agent(slug)
                self.assertTrue(body.strip(), f"{slug} 正文为空")

    def test_frontmatter_is_stripped(self):
        body = qc.read_agent("qc-conductor")
        self.assertFalse(body.startswith("---"), "frontmatter 未被剥掉")
        self.assertNotIn("description:", body.split("\n")[0])

    def test_missing_agent_raises(self):
        with self.assertRaisesRegex(RuntimeError, "找不到 Agent 定义"):
            qc.read_agent("qc-does-not-exist")

    def test_every_agent_is_registered_in_the_contract(self):
        contract = (
            Path(qc.HERE) / ".claude" / "rules" / "agent-authority-contract.md"
        ).read_text(encoding="utf-8")
        for slug in self.SLUGS:
            with self.subTest(slug=slug):
                self.assertIn(slug, contract, f"{slug} 未登记进责权利契约")


if __name__ == "__main__":
    unittest.main()
