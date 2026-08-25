#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
质检团队 · 全自动版
====================
一条命令，自动跑完:
  0. Crossref 真实核验文献 (不靠AI)
  1. 主审  = Claude Code  (claude 命令)
  2. 复核  = Codex        (codex 命令)
  3. 整理  = MiniMax      (API)
  → 输出完整质检报告

用法:
  python3 qc.py 稿件.docx
  python3 qc.py 稿件.txt

环境要求 (你的 Mac 已满足):
  - claude 命令可用
  - codex 命令可用
  - .env 里填了 MINIMAX_API_KEY
"""
import sys, os, subprocess, json, urllib.request, time, tempfile, shutil

HERE = os.path.dirname(os.path.abspath(__file__))
AGENTS = os.path.join(HERE, '.claude', 'agents')
REPORTS = os.path.join(HERE, 'reports')

# ---------- 读取 .env 里的密钥 (绝不写死在代码里) ----------
def load_env():
    env_path = os.path.join(HERE, '.env')
    cfg = {}
    if os.path.exists(env_path):
        for line in open(env_path, encoding='utf-8'):
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k, v = line.split('=', 1)
                cfg[k.strip()] = v.strip().strip('"').strip("'")
    # 也允许从系统环境变量读
    for k in ['MINIMAX_API_KEY', 'MINIMAX_API_URL', 'MINIMAX_MODEL']:
        if k not in cfg and os.environ.get(k):
            cfg[k] = os.environ[k]
    return cfg

def read(path):
    with open(path, encoding='utf-8', errors='ignore') as file:
        return file.read()

def read_agent(slug):
    """读取 .claude/agents/<slug>.md 并剥掉 YAML frontmatter，只把正文当角色提示词。

    Agent 定义同时被 Claude Code 的 Task 工具和本脚本使用，因此文件必须带
    frontmatter；frontmatter 是给运行时读的元数据，不该混进模型提示词。
    """
    path = os.path.join(AGENTS, f'{slug}.md')
    if not os.path.exists(path):
        raise RuntimeError(f'找不到 Agent 定义: {path}')
    text = read(path)
    if text.startswith('---'):
        parts = text.split('---', 2)
        if len(parts) == 3:
            return parts[2].lstrip('\n')
    return text

def read_docx(path):
    """读取 Word 正文和表格，不在原稿旁边生成中间 TXT。"""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            '读取 Word 需要 python-docx。请先运行: python3 -m pip install python-docx'
        ) from e

    try:
        document = Document(path)
    except Exception as e:
        raise RuntimeError(f'无法读取 Word 文件，请确认它是有效的 .docx 文件: {e}') from e

    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(' | '.join(cells))
    text = '\n'.join(lines).strip()
    if not text:
        raise RuntimeError('Word 文件中没有读取到正文或表格文字。扫描版 Word 请先进行 OCR。')
    return text

def load_manuscript(path):
    """自动识别 DOCX 或 UTF-8 文本稿件。"""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        print('  ✓ 已自动读取 Word 正文和表格，无需手动转 TXT')
        return read_docx(path)
    if ext in {'.txt', '.md'}:
        text = read(path).strip()
        if not text:
            raise RuntimeError('稿件内容为空。')
        return text
    raise RuntimeError('目前支持 .docx、.txt 和 .md 文件。旧版 .doc 请先另存为 .docx。')

def ensure_ready(cfg):
    """在产生模型费用前完成一次性配置检查。"""
    if not cfg.get('MINIMAX_API_KEY'):
        raise RuntimeError(
            '尚未配置 MiniMax。请把 .env.example 复制为 .env，并填写 MINIMAX_API_KEY。'
        )
    missing = [cmd for cmd in ('claude', 'codex') if not shutil.which(cmd)]
    if missing:
        raise RuntimeError(f"找不到命令: {', '.join(missing)}。请先安装并登录对应 CLI。")

def verify_references(doc_text):
    """将统一提取出的文本交给 Crossref 脚本，并在完成后删除临时文件。"""
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(
            mode='w', encoding='utf-8', suffix='.txt', prefix='qc-team-', delete=False
        ) as temp_file:
            temp_file.write(doc_text)
            temp_path = temp_file.name
        result = subprocess.run(
            [sys.executable, os.path.join(HERE, 'scripts', 'verify_refs.py'), temp_path],
            timeout=1800,
        )
        if result.returncode != 0:
            raise RuntimeError('Crossref 文献核验失败，流程已停止。')
    except subprocess.TimeoutExpired as e:
        raise RuntimeError('Crossref 文献核验超时，流程已停止。') from e
    finally:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)

# ---------- 第1步: 调 Claude Code 当主审 ----------
def call_claude(role, doc, refs_fact):
    prompt = f"{role}\n\n=== 文献核验结果(事实底座,以此为准) ===\n{refs_fact}\n\n=== 待质检稿件 ===\n{doc}\n\n请按你的角色输出主审意见。"
    print('  ▶ 主审(Claude Code)正在核查...(可能要一两分钟)')
    try:
        # claude -p 是非交互模式,直接给prompt拿输出
        r = subprocess.run(['claude', '-p', prompt],
                           capture_output=True, text=True, timeout=600,
                           encoding='utf-8', errors='replace')
        if r.returncode != 0:
            raise RuntimeError(f'Claude 主审调用失败: {r.stderr[:500]}')
        return r.stdout.strip()
    except FileNotFoundError:
        raise RuntimeError('找不到 claude 命令。请确认 Claude Code 已安装。')
    except subprocess.TimeoutExpired:
        raise RuntimeError('Claude 主审超时。稿件可能太长，可分段后重试。')

# ---------- 第2步: 调 Codex 当复核 ----------
def call_codex(role, doc, refs_fact, lead_review):
    prompt = f"{role}\n\n=== 文献核验结果 ===\n{refs_fact}\n\n=== 稿件 ===\n{doc}\n\n=== 主审意见(你要挑它的错) ===\n{lead_review}\n\n请按你的角色输出独立复核意见。"
    print('  ▶ 复核(Codex)正在独立核验、挑主审的错...')
    try:
        # codex exec 是非交互执行模式
        r = subprocess.run(['codex', 'exec', prompt],
                           capture_output=True, text=True, timeout=600,
                           encoding='utf-8', errors='replace')
        if r.returncode != 0:
            raise RuntimeError(f'Codex 复核调用失败: {r.stderr[:500]}')
        return r.stdout.strip()
    except FileNotFoundError:
        raise RuntimeError('找不到 codex 命令。')
    except subprocess.TimeoutExpired:
        raise RuntimeError('Codex 复核超时。')

# ---------- 第3步: 调 MiniMax API 当整理 ----------
def call_minimax(role, lead_review, cross_review, refs_fact, cfg):
    key = cfg.get('MINIMAX_API_KEY', '')
    url = cfg.get('MINIMAX_API_URL', 'https://api.minimaxi.com/v1/text/chatcompletion_v2')
    model = cfg.get('MINIMAX_MODEL', 'MiniMax-Text-01')
    if not key:
        raise RuntimeError('.env 里没有 MINIMAX_API_KEY。请填入后重跑。')

    user_content = (f"{role}\n\n=== 文献核验结果 ===\n{refs_fact}\n\n"
                    f"=== 主审意见 ===\n{lead_review}\n\n"
                    f"=== 复核意见 ===\n{cross_review}\n\n"
                    f"请综合两方,按你的角色输出一份包含全部8个板块的完整质检报告。")
    body = {
        "model": model,
        "messages": [{"role": "user", "content": user_content}],
    }
    print('  ▶ 整理(MiniMax)正在综合两方、推荐替代、出报告...')
    try:
        req = urllib.request.Request(url,
            data=json.dumps(body).encode('utf-8'),
            headers={'Content-Type': 'application/json',
                     'Authorization': f'Bearer {key}'})
        resp = json.loads(urllib.request.urlopen(req, timeout=600).read())
        # MiniMax 返回结构: choices[0].message.content
        return resp['choices'][0]['message']['content']
    except Exception as e:
        raise RuntimeError(
            f'MiniMax API 调用失败: {e}\n'
            '请核对 .env 里的 MINIMAX_API_URL、MINIMAX_MODEL 和 API Key。'
        ) from e

# ---------- 主流程 ----------
def main():
    if len(sys.argv) < 2:
        print('用法: python3 qc.py "/完整路径/稿件.docx"'); sys.exit(1)
    docpath = os.path.abspath(os.path.expanduser(sys.argv[1]))
    if not os.path.exists(docpath):
        print(f'找不到文件: {docpath}'); sys.exit(1)

    try:
        cfg = load_env()
        ensure_ready(cfg)
        doc = load_manuscript(docpath)
        name = os.path.splitext(os.path.basename(docpath))[0]
        stamp = time.strftime('%Y-%m-%d_%H%M%S')
        os.makedirs(REPORTS, exist_ok=True)

        print('='*50)
        print(f'  三角色质检团队启动 · 稿件: {name}')
        print('  Claude 主审 → Codex 复核 → MiniMax 整理')
        print('='*50)

        print('\n▶ 第0步: Crossref 真实核验文献...')
        verify_references(doc)
        refs_path = os.path.join(REPORTS, '_refs_verified.json')
        refs_fact = read(refs_path) if os.path.exists(refs_path) else '(无 DOI 或核验未产出)'

        print('\n▶ 第1步: Claude 主审')
        lead = call_claude(read_agent('qc-conductor'), doc, refs_fact)
        open(os.path.join(REPORTS, f'{name}_1主审.md'), 'w', encoding='utf-8').write(lead)

        print('\n▶ 第2步: Codex 复核')
        cross = call_codex(read_agent('qc-verifier'), doc, refs_fact, lead)
        open(os.path.join(REPORTS, f'{name}_2复核.md'), 'w', encoding='utf-8').write(cross)

        print('\n▶ 第3步: MiniMax 整理')
        final = call_minimax(read_agent('qc-reporter'), lead, cross, refs_fact, cfg)
        outpath = os.path.join(REPORTS, f'{name}_质检报告_{stamp}.md')
        open(outpath, 'w', encoding='utf-8').write(final)

        print('\n' + '='*50)
        print('  ✅ 三个角色已全部完成，最终报告:')
        print(f'  {outpath}')
        print('='*50)
    except RuntimeError as e:
        print(f'\n❌ {e}', file=sys.stderr)
        print('流程已停止，不会把不完整结果显示为成功。', file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()
