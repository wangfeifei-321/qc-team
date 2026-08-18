#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
把 .docx 稿件转成纯文本，供质检使用。
用法: python3 docx2txt.py 稿件.docx
输出: 稿件.txt (同目录)
需要: pip install python-docx
"""
import sys, os
try:
    from docx import Document
except ImportError:
    print('请先安装: pip install python-docx --break-system-packages')
    sys.exit(1)

if len(sys.argv) < 2:
    print('用法: python3 docx2txt.py 稿件.docx')
    sys.exit(1)

path = sys.argv[1]
doc = Document(path)
lines = []
for p in doc.paragraphs:
    if p.text.strip():
        lines.append(p.text)
# 表格里的引用/数据也要带上
for tb in doc.tables:
    for row in tb.rows:
        cells = [c.text.strip() for c in row.cells if c.text.strip()]
        if cells:
            lines.append(' | '.join(cells))

out = os.path.splitext(path)[0] + '.txt'
open(out, 'w', encoding='utf-8').write('\n'.join(lines))
print(f'✅ 已转换 → {out}  ({len(lines)} 段)')
